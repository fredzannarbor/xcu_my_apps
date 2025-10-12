#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Tue Jan  5 14:51:25 2021

@author: fred
"""
import os

from docx import Document
from docxcompose.composer import Composer

directory = '/Users/fred/Dropbox/Documents/nimble/editorial/GPT3 Society'

master = Document()
composer = Composer(master)

for filename in sorted(os.listdir(directory)):
    if filename.endswith(".docx"): 
        f = os.path.join(directory, filename)
        doc1 = Document(f)
        p = doc1.paragraphs[0]
        p.insert_paragraph_before(filename, style='Heading1')
        doc1.save(filename)
        composer.append(doc1)
        #print('adding paragraph', filename)
        print('appending', f)

        continue
    else:
        continue
    

composer.save(os.path.join(directory, "new2combined.docx"))