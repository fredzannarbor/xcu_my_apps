import datetime
import json
import re
from pathlib import Path

import pandas as pd
import requests

try:
    from pypandoc import convert_file
except ImportError as e:
    print("Pandoc not installed. No pypandoc module")
    print(e)

import streamlit as st
from docx import Document
# import Inches
from docx.shared import Inches


# this function
def metadatas2docx(metadatas, thisdoc_dir, target_df="LSI"):
    # create a dataframe from the metadatas
    print(">entering metadatas2df")
    # print('metadata keys', metadatas.keys())
    # apply regex to the values in the metadatas dict
    for k, v in metadatas.items():
        if isinstance(v, str):
            metadatas[k] = re.sub(r'^\n+', '', v)


    if metadatas['mode'] == "assess":
        return
    if target_df == "LSI":
        # initializing variables

        ISBN = str(metadatas['ISBN'])
        metadatas["AI_byline"] = "Enhanced by Nimble Books AI"
        copyright_credit_pd = "(c) " + str(datetime.datetime.now().year) + " " + metadatas["publisher"]
        metadatas['foreword_sig'] = "Cincinnatus [AI]" #replace with varaib le
        foreword_sig_line = "--" + metadatas['foreword_sig']

        tldr1_text = metadatas["TLDR (one word)"][1]
        tldr_vanilla_text = metadatas["TLDR"][1]
        ELI5_text = metadatas["ELI5"][1]

        title = {"string": metadatas['title'], "style": "Title Page Title"}
        if metadatas["subtitle"]:
            subtitle = {"string": metadatas["subtitle"], "style": "Title Page Subtitle"}
        else:
            subtitle = {"string": "TK", "style": "Title Page Subtitle"}

        byline = {"string": metadatas["author"], "style": "Title Page Byline"}

        AI_byline = {"string": metadatas["AI_byline"], "style": "Title Page Byline"}
        imprint = {"string": metadatas["imprint"], "style": "Title Page Nimble Logotype"}
        pagebreak = {"string": "Page Break", "style": "Page Break"}
        publishing_information = {"string": "Publishing Information", "style": "Heading 2"}
        copyright_line = {"string": copyright_credit_pd, "style": "Body Text - flush left"}
        if ISBN:
            ISBN_line = {"string": "ISBN: " + str(metadatas['ISBN']), "style": "Body Text - flush left"}
        else:
            metadatas["ISBN"] = "TK"
            ISBN_text = "ISBN: " + metadatas['ISBN']
            ISBN_line = {"string": ISBN_text, "style": "Body Text - flush left"}
        series_name = "AI Lab for Book-Lovers "
        series_number = "No. " #metadatas["series_number"]
        series_name_line = {"string": series_name, "style": "Body Text - flush left"}
        series_number_line = {"string": series_number, "style": "Body Text - flush left"}
        series_info_text = "Humans and AI making books richer, more diverse, and more surprising."
        series_info_line = {"string": series_info_text, "style": "Body Text - flush left"}
        bibliographic_kw = {"string": "Bibliographic Keywords", "style": "Heading 2"}
        if iter(metadatas['Publisher-supplied Keywords']) and len(metadatas['Publisher-supplied Keywords']) > 0:
            author_supplied_kw_string = metadatas['Publisher-supplied Keywords']
            author_supplied_kw_heading = {"string": "Publisher-supplied Keywords", "style": "Heading 2"}
            author_supplied_kw_line = {"string": author_supplied_kw_string, "style": "Heading 2"}
        else:
            author_supplied_kw_heading = {"string": "", "style": "Heading 1"}
            author_supplied_kw_line = {"string": "", "style": "Body Text"}
        algorithmically_generated_kw_line = {"string": "AI-generated Keyword Phrases", "style": "Heading 2"}
        algorithmically_generated_kw = {"string": metadatas["Bibliographic Keyword Phrases"][1], "style": "Body Text - flush left"}

        pagebreak3 = {"string": "Page Break", "style": "Page Break"}
        foreword_line = {"string": "Foreword", "style": "Heading 1"}
        foreword = {"string": metadatas["Foreword"], "style": "Body Text"}
        foreword_sig = {"string": metadatas["foreword_sig"], "style": "Signature"}
        abstracts = {"string": "Abstracts", "style": "Heading 1"}
        scientific_style = {"string": "Scientific Style", "style": "Heading 2"}
        scientific_style_text = {"string": metadatas["Scientific Style"][1], "style": "Body Text"}

        tldr1_heading = {"string": "TL;DR (one word)", "style": "Heading 2"}
        tldr1_line = {"string": tldr1_text, "style": "Body Text"}
        tldr_vanilla = {"string": "TL;DR (vanilla)", "style": "Heading 2"}
        tldr_vanilla_line = {"string": tldr_vanilla_text, "style": "Body Text"}
        ELI5 = {"string": "Explain It To Me Like I'm Five Years Old", "style": "Heading 2"}
        ELI5_line = {"string": ELI5_text, "style": "Body Text"}
        pagebreak5 = {"string": "Page Break", "style": "Page Break"}
        viewpoints_heading = {"string": "Viewpoints", "style": "Heading 1"}
        viewpoints_motivation = {"string": "These perspectives increase the reader's exposure to viewpoint diversity.", "style": "Body Text"}
        MAGA_heading = {"string": "MAGA Perspective", "style": "Heading 2"}
        MAGA_text = {"string": metadatas["Hostile MAGA Perspective"][1], "style": "Body Text"}
        FormalDissent_heading = {"string": "Formal Dissent", "style": "Heading 2"}
        FormalDissent_text = {"string": metadatas["Formal Dissent"][1], "style": "Body Text"}
        RedTeamCritique_heading = {"string": "Red Team Critique", "style": "Heading 2"}
        RedTeamCritique_text = {"string": metadatas["Red Team Critique"][1], "style": "Body Text"}
        ActionItems_heading = {"string": "Action Items", "style": "Heading 2"}
        ActionItems_text = {"string": metadatas["Suggested Action Items"][1], "style": "Body Text"}
        pagebreak4 = {"string": "Page Break", "style": "Page Break"}
        recursive_headings = {"string": "Summaries", "style": "Heading 1"}
        recursive_methods_heading = {"string": "Methods", "style": "Heading 2"}
        recursive_methods_text = {"string": metadatas["Methods"], "style": "Body Text"}
        extractive_summary_heading = {"string": "Extractive Summary", "style": "Heading 2"}
        extractive_summary_signature = {
            "string": "Most representative sentences identified by sumy implementation of LexRank algorithm.",
            "style": "Signature"}
        extractive_summary_text = {"string": metadatas["extractive_summary"], "style": "Body Text"}
        extractive_synopsis_heading = {"string": "Extractive Synopsis", "style": "Heading 2"}
        extractive_synopsis_signature = {"string": "Most representative several sentences identified by sumy implementation of LexRank algorithm.", "style": "Signature"}
        extractive_synopsis_text = {"string": metadatas["extractive_synopsis"], "style": "Body Text"}
        # add all dicts to a list
        rows = [title, subtitle, byline, AI_byline, imprint, pagebreak, publishing_information, copyright_line,
                ISBN_line, series_name_line, series_number_line, series_info_line,
                bibliographic_kw, author_supplied_kw_heading, author_supplied_kw_line, algorithmically_generated_kw_line,
                algorithmically_generated_kw, pagebreak3, foreword_line, foreword, foreword_sig, abstracts,
                scientific_style,
                scientific_style_text, tldr1_heading, tldr1_line,
                tldr_vanilla, tldr_vanilla_line, ELI5, ELI5_line, ActionItems_heading, ActionItems_text, pagebreak5, viewpoints_heading, viewpoints_motivation, FormalDissent_heading,
                FormalDissent_text, RedTeamCritique_heading, RedTeamCritique_text, MAGA_heading, MAGA_text,
                pagebreak4, recursive_headings,
                recursive_methods_heading,
                recursive_methods_text]
        recursive_rows = []

        for index, r in enumerate(metadatas["Recursive Summaries"]):
            recursive_summaries_heading = {"string": "Recursive Summary Round " + str(index), "style": "Heading 2"}
            recursive_rows.append(recursive_summaries_heading)
            recursive_summaries_text = {"string": r, "style": "Body Text"}
            recursive_rows.append(recursive_summaries_text)
        rows.extend(recursive_rows)

        if metadatas["page_by_page_results"]:
            pagebypage_rows = []
            pagebypage_page_heading = {"string": "Page by Page Summaries ", "style": "Heading 2"}
            pagebypage_rows.append(pagebypage_page_heading)
            pagesummaries = metadatas["page_by_page_results"]

            for pagepair in pagesummaries:
                pagesubhead =  "Page " + str(pagepair[0])
                pagetext_pp = re.sub(r'^\n\n', '', pagepair[1])
                pageheading = {"string": pagesubhead, "style": "Heading 3"}
                pagetext = {"string": pagetext_pp, "style": "Body Text"}
                pagebypage_rows.append(pageheading)
                pagebypage_rows.append(pagetext)
            rows.extend(pagebypage_rows)

        if metadatas['Text2CoverImage Prompt']:
            imageprompt_rows = []
            pagebreak5 = {"string": "Page Break", "style": "Page Break"}
            imageprompt_rows.append(pagebreak5)
            imageprompt_heading = {"string": "Moods", "style": "Heading 1"}
            imageprompt_moods_explanation = {"string": "Multimodal generative AI is used to morph the informational and emotional content of this publication into visual expression. Highly experimental.", "style": "Body Text"}
            imageprompt_rows.append(imageprompt_heading)
            imageprompt_rows.append(imageprompt_moods_explanation)
            image_placeholder = {"string": "Image Placeholder", "style": "Image"}
            imageprompt_rows.append(image_placeholder)
            imageprompt_caption = f"The AI-generated prompt was: {metadatas['Text2CoverImage Prompt'][1]}"
            imageprompt_text = {"string": imageprompt_caption, "style": "Caption"}
            imageprompt_rows.append(imageprompt_text)
            #rows.extend(imageprompt_rows)

        if metadatas['Text2CoverImage Prompt (Stable Diffusion)']:
            imageprompt_rows = []
            imageprompt_heading = {"string": "Stable Diffusion Moods", "style": "Heading 2"}
            #imageprompt_rows.append(imageprompt_heading)
            image_placeholder = {"string": "Image Placeholder", "style": "Image"}
            imageprompt_rows.append(image_placeholder)
            imageprompt_caption = f"The AI-generated prompt was: {metadatas['Text2CoverImage Prompt (Stable Diffusion)'][1]}"
            imageprompt_text = {"string": imageprompt_caption, "style": "Caption"}
            imageprompt_rows.append(imageprompt_text)
            rows.extend(imageprompt_rows)
        print('made it to T2MP_response')
        if metadatas['Text2MoodImagePrompt_response'] is not None:
            # get images`
            imageprompt_rows = []
            imageprompt_heading = {"string": "DALLE2 Moods", "style": "Heading 2"}
            #imageprompt_rows.append(imageprompt_heading)
            count = 0
            d = metadatas['Text2MoodImagePrompt_response']
            print(d)
            for item in d["data"]:
                url = item["url"]
                r = requests.get(url)
                image_file_path = f"{thisdoc_dir}/DALLE2_images/{count}.jpg"
                with open(image_file_path, "wb") as f:
                    f.write(r.content)
                count += 1
                image_placeholder = {"string": image_file_path, "style": "Image"}
                imageprompt_rows.append(image_placeholder)
                imageprompt_caption = f"The AI-generated prompt was: {metadatas['Text2MoodImagePrompt'][1]}"
                imageprompt_text = {"string": imageprompt_caption, "style": "Caption"}
                imageprompt_rows.append(imageprompt_text)
            rows.extend(imageprompt_rows)

        if metadatas['midjourney_response']:
            midjourney_rows = []
            midjourney_heading = {"string": "Midjourney Mood", "style": "Heading 2"}
            midjourney_rows.append(midjourney_heading)
            # get image
            midjourney_image = requests.get(metadatas['midjourney_response'])
            midjourney_image_path = thisdoc_dir + "/midjourney_image.jpg"
            with open(midjourney_image_path, 'wb') as f:
                f.write(midjourney_image.content)
            # insert image placeholder
            midjourney_image_placeholder = {"string": midjourney_image_path, "style": "Image"}
            midjourney_rows.append(midjourney_image_placeholder)
            midjourney_caption = f"The AI-generated instruction to the AI was as follows: {metadatas['midjourney prompt']}"
            midjourney_caption_text = {"string": midjourney_caption, "style": "Caption"}
            midjourney_rows.append(midjourney_caption_text)
            rows.extend(midjourney_rows)

        with open(thisdoc_dir + '/rows.json', "w") as f:
            json.dump(rows, f)

        # create a dataframe from a list of strings

        docx_production_df = pd.DataFrame(rows)
        docx_production_df.to_json(thisdoc_dir + '/' + "target_df.json")

        document = Document("resources/docx/FrontMatter2.docx")
        try:
            for index, row in docx_production_df.iterrows():
                print(row)
                if row["style"] == "Page Break":
                    document.add_page_break()
                    continue
                if row["string"] == "":
                    continue

                # remove newline if it occurs at beginning of string
                row["string"] = re.sub(r'^\n', '', row["string"])
                # if a newline is surrouonded only by spaces,
                # remove the spaces
                row["string"] = re.sub(r'\n\s*\n', '\n\n', row["string"])
                # if two or more newlines are found in a row, convert them to only one
                row["string"] = re.sub(r'\n{2,}', '\n', row["string"])
                # if a newline is found at the end of a string, remove it
                row["string"] = re.sub(r'\n$', '', row["string"])

                if row["style"] == "Image":
                    if row["string"] != "Image Placeholder":
                        try:
                            document.add_picture(row["string"], width=Inches(5.5))
                        except Exception as e:
                            errmessage = str(e) + 'error in adding image'
                            print(errmessage)
                            document.add_paragraph("Image not found, check log to see if it was generated.", row["style"])
                    else:
                        document.add_paragraph(row["string"], row["style"])
                else:
                    # count newlines in the text-bearing string
                    newline_count = row["string"].count('\n')
                # now if there are multiple paragraphs in the string, split them and add them as separate paragraphs
                    if newline_count >= 1:
                        paragraphs = row["string"].split('\n')
                        for p in paragraphs:
                            q = "Problem is here"
                            document.add_paragraph(p, row["style"])
                    else:
                        document.add_paragraph(row["string"], row["style"])

        except Exception as e:
            errmessage = str(e) + 'error in creating word doc'
            print(errmessage)


        document.save(thisdoc_dir + '/' + "frontmatter.docx")

        # convert the doc to pdf
        try:
            convert_file(thisdoc_dir + '/' + "frontmatter.docx", 'pdf',
                              outputfile=thisdoc_dir + '/' + "frontmatter.pdf")
        except Exception as e:
            errmessage = str(e) + 'error in creating interior postscript'
            print(errmessage)

        return docx_production_df, document


def metadatas2bookjson(metadatas, thisdoc_dir, distributor="LSI"):
    # convert the metadata to a json file
    bookjson = {}
    bookjson["BookTitle"] = metadatas["title"]
    bookjson["SubTitle"] = metadatas["subtitle"]
    bookjson["Byline"] = metadatas["author"]
    bookjson["ImprintText"] = metadatas["imprint"]
    bookjson["ImageFileName"] = ""  # metadatas["cover_image"]
    bookjson["settings"] = "default"
    bookjson["distributor"] = "LSI"
    bookjson["InvertedColor"] = "White"
    bookjson["DominantColor"] = "Nimble Maroon"
    bookjson["BaseFont"] = "CMU Serif Roman"
    bookjson["trimsizewidth"] = 8.5  # metadatas["trim_size"][0]
    bookjson["trimsizeheight"] = 11  # metadatas["trim_size"][1]
    bookjson["spinewidth"] = metadatas["spinewidth"]
    print(metadatas["Recursive Summaries"][-1])
    print(metadatas["Publisher-supplied synopsis"])
    backtext = metadatas["Recursive Summaries"][-1] + metadatas["Publisher-supplied synopsis"]
    bookjson["backtext"] = metadatas["Recursive Summaries"][-1]
    st.write(bookjson)
    with open(thisdoc_dir + '/' + Path(thisdoc_dir).stem + "_book.json", 'w') as f:
        json.dump(bookjson, f)

    with open("bookjsonfiles/" + Path(thisdoc_dir).stem + "_book.json", 'w') as f:
        json.dump(bookjson, f)

    return bookjson


def metadatas2internationaledition(metadatas, thisdoc_dir, languages, international_presets, edition_name):
    # convert

    # add bespoke content that motivatess the international edition

    # translate selected metadata to other languages
    international_text = 'placeholder' # hgrtext2googlecloudtranslate.py(metadatas, languages, international_presets)
    # assemble supplementary front matter section

    return "successfully built book json file for Scribus"
