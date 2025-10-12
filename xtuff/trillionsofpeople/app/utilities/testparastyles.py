""" test document and para styles"""
from docx import Document
from docx import __version__;
from docx.enum.style import WD_STYLE_TYPE

print(__version__)
docx_filename = "test/docx/syz2.docx"
doc = Document(docx_filename)
# print(dir(doc))
# print(dir(doc.styles))

# print('------------------------------')

paragraphs = doc.paragraphs
print(dir(paragraphs.remove))
styles = doc.styles
print(styles['Default Paragraph Font'].font.name)
print('------------------------------')
paragraph_styles = [s for s in styles if s.type == WD_STYLE_TYPE.PARAGRAPH]
print('paragraph styles: {}'.format(paragraph_styles))
print('------------------------------')
if "Times New Roman" in [s.font.name for s in paragraph_styles]:
    print("Times New Roman is in the paragraph styles")
character_styles = [s for s in styles if s.type == WD_STYLE_TYPE.CHARACTER]
print(paragraph_styles)
for c in character_styles:
    print(c.name)
    print(c.font)
    print('------------------------------')
if "Tim    ('cccc      es New Roman" in [s.font.name for s in character_styles]:
    print("Times New Roman is in the character styles")
print(character_styles)

print(dir(styles))
#fonts = doc.fonts
print(len(paragraphs))

for p in doc.paragraphs:

    print(p.style.font.name)
    print(p.style.font.size)
    p.paragraph_format.line_spacing=None
    runs = p.runs
    for r in runs:
        r.style = None
        r.font.name = None
        print(r.style.name)
try:
    doc.save('test/docx/syz_bodytext.docx')
except Exception as e:
    print(e)