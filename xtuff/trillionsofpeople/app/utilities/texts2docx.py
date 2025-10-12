"""input list of text objects or files, output docx"""
import argparse
import os

from docx import Document


def texts2docx(text_list, docx_filename, output_dir, working_dir, basedocx=None):
    if basedocx:
        document = Document(basedocx)
    else:
        document = Document()
        try:
            for text in text_list:  # rint(text)
                document.add_paragraph(text)
            #print(text_list)
        except Exception as e:
            print(e, 'error in texts2docx')
        document.save(os.path.join(output_dir, docx_filename))
        
    return docx_filename
        
def add_to_document(document, text):

    document = Document(document)# add text to document
    
    for paragraph in text:
        document.add_paragraph(paragraph, style='Body Text')
    document.add_page_break()
    return document

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description='input list of text objects or files, output docx')
    parser.add_argument("--basedocx", help = "base Word file", default = "/Users/fred/bin/nimble/dvids/templates/cool_ships2.docx")
    parser.add_argument("--output_dir", help = "output directory", default = "/Users/fred/bin/nimble/unity/output/")
    parser.add_argument("--title", help = "title", default = "Front_Matter")
    parser.add_argument("--keywords", help = "keywords", default = "cool, ships")
    parser.add_argument("--text_list", help = "list of text objects", default = "[text, title, abstract, body]")

    args = parser.parse_args()
    
    basedocx = args.basedocx
    output_dir = args.output_dir

    title = args.title

    keywords = args.keywords

    text_list = args.text_list
    text_list = text_list.split(',')

    text_list = [title, keywords, abstract, body]
    texts2docx(text_list, 'test.docx', output_dir, '.')