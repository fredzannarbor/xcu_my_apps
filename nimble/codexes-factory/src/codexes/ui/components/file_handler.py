"""
File Handler Component

Handles file uploads, format extraction, and directory browsing for the Stage-Agnostic UI.
"""

import os
import zipfile
import tempfile
from pathlib import Path
from typing import List, Tuple, Dict, Any, Optional
import streamlit as st
from io import BytesIO, StringIO
try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False


class FileHandler:
    """Handles file operations for the universal content input."""
    
    def __init__(self):
        # Base formats that are always supported
        self.supported_formats = ['.txt', '.md', '.rtf']
        
        # Add optional formats if libraries are available
        if DOCX_AVAILABLE:
            self.supported_formats.append('.docx')
        if PDF_AVAILABLE:
            self.supported_formats.append('.pdf')
        
        self.allowed_directories = ['data/', 'output/']
        self.max_file_size = 50 * 1024 * 1024  # 50MB limit
    
    def extract_content(self, uploaded_file) -> str:
        """Extract text content from uploaded file."""
        try:
            file_extension = Path(uploaded_file.name).suffix.lower()
            
            if file_extension == '.txt':
                return self._extract_txt(uploaded_file)
            elif file_extension == '.md':
                return self._extract_markdown(uploaded_file)
            elif file_extension == '.docx':
                return self._extract_docx(uploaded_file)
            elif file_extension == '.pdf':
                return self._extract_pdf(uploaded_file)
            elif file_extension == '.rtf':
                return self._extract_rtf(uploaded_file)
            else:
                st.warning(f"Unsupported file format: {file_extension}")
                return ""
        
        except Exception as e:
            st.error(f"Error extracting content from {uploaded_file.name}: {str(e)}")
            return ""
    
    def _extract_txt(self, uploaded_file) -> str:
        """Extract content from text file."""
        try:
            # Try UTF-8 first, then fallback to other encodings
            content = uploaded_file.read()
            if isinstance(content, bytes):
                try:
                    return content.decode('utf-8')
                except UnicodeDecodeError:
                    try:
                        return content.decode('latin-1')
                    except UnicodeDecodeError:
                        return content.decode('utf-8', errors='ignore')
            return str(content)
        except Exception as e:
            st.error(f"Error reading text file: {str(e)}")
            return ""
    
    def _extract_markdown(self, uploaded_file) -> str:
        """Extract content from markdown file."""
        return self._extract_txt(uploaded_file)  # Markdown is just text
    
    def _extract_docx(self, uploaded_file) -> str:
        """Extract content from Word document."""
        if not DOCX_AVAILABLE:
            st.warning("python-docx not available. Cannot process .docx files.")
            return ""
        
        try:
            doc = docx.Document(uploaded_file)
            content = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content.append(paragraph.text)
            
            return '\n\n'.join(content)
        
        except Exception as e:
            st.error(f"Error reading Word document: {str(e)}")
            return ""
    
    def _extract_pdf(self, uploaded_file) -> str:
        """Extract content from PDF file."""
        if not PDF_AVAILABLE:
            st.warning("PyPDF2 not available. Cannot process .pdf files.")
            return ""
        
        try:
            pdf_reader = PyPDF2.PdfReader(uploaded_file)
            content = []
            
            for page in pdf_reader.pages:
                text = page.extract_text()
                if text.strip():
                    content.append(text)
            
            return '\n\n'.join(content)
        
        except Exception as e:
            st.error(f"Error reading PDF: {str(e)}")
            return ""
    
    def _extract_rtf(self, uploaded_file) -> str:
        """Extract content from RTF file (basic implementation)."""
        try:
            # Basic RTF parsing - just extract text between braces
            content = uploaded_file.read()
            if isinstance(content, bytes):
                content = content.decode('utf-8', errors='ignore')
            
            # Very basic RTF parsing - remove control words
            import re
            # Remove RTF control words
            content = re.sub(r'\\[a-z]+\d*', '', content)
            # Remove braces
            content = re.sub(r'[{}]', '', content)
            # Clean up whitespace
            content = re.sub(r'\s+', ' ', content).strip()
            
            return content
        
        except Exception as e:
            st.error(f"Error reading RTF file: {str(e)}")
            return ""
    
    def extract_zip_contents(self, uploaded_zip) -> List[Tuple[str, str]]:
        """Extract contents from ZIP archive."""
        content_items = []
        
        try:
            with zipfile.ZipFile(uploaded_zip, 'r') as zip_ref:
                file_list = zip_ref.namelist()
                
                # Filter for supported files
                supported_files = [
                    f for f in file_list 
                    if Path(f).suffix.lower() in self.supported_formats
                    and not f.startswith('__MACOSX/')  # Skip Mac metadata
                    and not f.startswith('.')  # Skip hidden files
                ]
                
                if not supported_files:
                    st.warning("No supported files found in ZIP archive")
                    return content_items
                
                # Limit number of files to prevent overwhelming
                if len(supported_files) > 50:
                    st.warning(f"ZIP contains {len(supported_files)} files. Processing first 50.")
                    supported_files = supported_files[:50]
                
                for file_path in supported_files:
                    try:
                        with zip_ref.open(file_path) as file_in_zip:
                            # Create a temporary file-like object
                            file_content = file_in_zip.read()
                            
                            # Create a mock uploaded file object
                            mock_file = type('MockFile', (), {
                                'name': Path(file_path).name,
                                'read': lambda: file_content
                            })()
                            
                            extracted_content = self.extract_content(mock_file)
                            if extracted_content.strip():
                                content_items.append((file_path, extracted_content))
                    
                    except Exception as e:
                        st.warning(f"Could not extract {file_path}: {str(e)}")
                        continue
        
        except Exception as e:
            st.error(f"Error processing ZIP file: {str(e)}")
        
        return content_items
    
    def list_directory_files(self, directory: str, supported_formats: List[str]) -> List[str]:
        """List files in a directory that match supported formats."""
        try:
            if directory not in self.allowed_directories:
                if hasattr(st, 'error'):
                    st.error(f"Directory {directory} is not allowed")
                else:
                    print(f"Directory {directory} is not allowed")
                return []
            
            directory_path = Path(directory)
            if not directory_path.exists():
                if hasattr(st, 'warning'):
                    st.warning(f"Directory {directory} does not exist")
                else:
                    print(f"Directory {directory} does not exist")
                return []
            
            files = []
            
            for file_path in directory_path.rglob('*'):
                if (file_path.is_file() and 
                    file_path.suffix.lower() in supported_formats and
                    not file_path.name.startswith('.')):  # Skip hidden files
                    # Use the path as-is since it's already relative to the directory
                    files.append(str(file_path))
            
            return sorted(files)
        
        except Exception as e:
            if hasattr(st, 'error'):
                st.error(f"Error listing directory {directory}: {str(e)}")
            else:
                print(f"Error listing directory {directory}: {str(e)}")
            return []
    
    def read_file(self, file_path: str) -> str:
        """Read content from a file path."""
        try:
            file_path_obj = Path(file_path)
            
            if not file_path_obj.exists():
                st.error(f"File {file_path} does not exist")
                return ""
            
            # Check if file is in allowed directory
            allowed = False
            for allowed_dir in self.allowed_directories:
                if file_path.startswith(allowed_dir):
                    allowed = True
                    break
            
            if not allowed:
                st.error(f"File {file_path} is not in an allowed directory")
                return ""
            
            # Read based on file extension
            file_extension = file_path_obj.suffix.lower()
            
            if file_extension in ['.txt', '.md']:
                with open(file_path_obj, 'r', encoding='utf-8') as f:
                    return f.read()
            
            elif file_extension == '.docx':
                if not DOCX_AVAILABLE:
                    st.warning("python-docx not available. Cannot process .docx files.")
                    return ""
                doc = docx.Document(file_path_obj)
                content = []
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        content.append(paragraph.text)
                return '\n\n'.join(content)
            
            elif file_extension == '.pdf':
                if not PDF_AVAILABLE:
                    st.warning("PyPDF2 not available. Cannot process .pdf files.")
                    return ""
                with open(file_path_obj, 'rb') as f:
                    pdf_reader = PyPDF2.PdfReader(f)
                    content = []
                    for page in pdf_reader.pages:
                        text = page.extract_text()
                        if text.strip():
                            content.append(text)
                    return '\n\n'.join(content)
            
            else:
                st.warning(f"Unsupported file format for reading: {file_extension}")
                return ""
        
        except Exception as e:
            st.error(f"Error reading file {file_path}: {str(e)}")
            return ""
    
    def get_file_info(self, file_path: str) -> Dict[str, Any]:
        """Get information about a file."""
        try:
            file_path_obj = Path(file_path)
            
            if not file_path_obj.exists():
                return {}
            
            stat = file_path_obj.stat()
            
            return {
                'name': file_path_obj.name,
                'size': stat.st_size,
                'size_human': self._format_file_size(stat.st_size),
                'extension': file_path_obj.suffix.lower(),
                'modified': stat.st_mtime
            }
        
        except Exception as e:
            st.error(f"Error getting file info for {file_path}: {str(e)}")
            return {}
    
    def _format_file_size(self, size_bytes: int) -> str:
        """Format file size in human readable format."""
        if size_bytes == 0:
            return "0 B"
        
        size_names = ["B", "KB", "MB", "GB"]
        i = 0
        while size_bytes >= 1024 and i < len(size_names) - 1:
            size_bytes /= 1024.0
            i += 1
        
        return f"{size_bytes:.1f} {size_names[i]}"
    
    def validate_file(self, uploaded_file) -> Tuple[bool, str]:
        """Validate uploaded file."""
        if not uploaded_file:
            return False, "No file provided"
        
        # Check file size
        if hasattr(uploaded_file, 'size') and uploaded_file.size > self.max_file_size:
            return False, f"File too large. Maximum size: {self._format_file_size(self.max_file_size)}"
        
        # Check file extension
        file_extension = Path(uploaded_file.name).suffix.lower()
        if file_extension not in self.supported_formats and file_extension != '.zip':
            return False, f"Unsupported file format: {file_extension}"
        
        return True, "File is valid"