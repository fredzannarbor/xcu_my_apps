"""
Supplementary Materials Loader

Loads and processes supplementary materials (txt, markdown, docx, pdf) for imprint configs.
Efficiently compacts content for inclusion in LLM calls with large context windows.
"""

import logging
from typing import Dict, List, Optional, Any
from pathlib import Path
import json

logger = logging.getLogger(__name__)


class SupplementaryMaterialsLoader:
    """Loads and processes supplementary materials from various file formats."""

    SUPPORTED_FORMATS = {'.txt', '.md', '.markdown', '.docx', '.pdf'}

    def __init__(self):
        """Initialize the loader."""
        self.loaded_materials = {}

    def load_materials(self, materials_dict: Dict[str, str]) -> Dict[str, str]:
        """
        Load supplementary materials from file paths.

        Args:
            materials_dict: Dictionary mapping titles to file paths

        Returns:
            Dictionary mapping titles to extracted text content
        """
        loaded = {}

        for title, file_path in materials_dict.items():
            try:
                content = self._load_single_file(file_path)
                if content:
                    loaded[title] = content
                    logger.info(f"Loaded supplementary material: {title} ({len(content)} chars)")
                else:
                    logger.warning(f"No content extracted from {file_path}")
            except Exception as e:
                logger.error(f"Failed to load {title} from {file_path}: {e}")

        self.loaded_materials = loaded
        return loaded

    def _load_single_file(self, file_path: str) -> Optional[str]:
        """
        Load a single file and extract text content.

        Args:
            file_path: Path to the file

        Returns:
            Extracted text content or None
        """
        path = Path(file_path)

        if not path.exists():
            logger.error(f"File not found: {file_path}")
            return None

        suffix = path.suffix.lower()

        if suffix not in self.SUPPORTED_FORMATS:
            logger.error(f"Unsupported file format: {suffix}")
            return None

        try:
            if suffix in {'.txt', '.md', '.markdown'}:
                return self._load_text_file(path)
            elif suffix == '.docx':
                return self._load_docx(path)
            elif suffix == '.pdf':
                return self._load_pdf(path)
        except Exception as e:
            logger.error(f"Error loading {file_path}: {e}")
            return None

    def _load_text_file(self, path: Path) -> str:
        """Load plain text or markdown file."""
        with open(path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()

    def _load_docx(self, path: Path) -> str:
        """Load Microsoft Word .docx file."""
        try:
            from docx import Document
        except ImportError:
            logger.error("python-docx not installed. Install with: uv add python-docx")
            return ""

        doc = Document(path)

        # Extract all paragraphs
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        # Extract text from tables
        table_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                table_text.append(" | ".join(row_text))

        content_parts = paragraphs
        if table_text:
            content_parts.append("\n\nTABLES:\n" + "\n".join(table_text))

        return "\n\n".join(content_parts)

    def _load_pdf(self, path: Path) -> str:
        """Load PDF file using PyMuPDF (fitz)."""
        try:
            import fitz  # PyMuPDF
        except ImportError:
            logger.error("PyMuPDF not installed. Install with: uv add pymupdf")
            return ""

        doc = fitz.open(path)
        text_parts = []

        for page_num, page in enumerate(doc, 1):
            text = page.get_text()
            if text.strip():
                text_parts.append(f"--- Page {page_num} ---\n{text}")

        doc.close()
        return "\n\n".join(text_parts)

    def compact_materials(
        self,
        materials: Dict[str, str],
        max_chars: Optional[int] = None,
        compression_strategy: str = "truncate"
    ) -> str:
        """
        Compact multiple materials into a single string for LLM context.

        Args:
            materials: Dictionary of title -> content
            max_chars: Maximum characters (None = no limit)
            compression_strategy: "truncate", "summarize", or "smart"

        Returns:
            Compacted string suitable for LLM prompt
        """
        if not materials:
            return ""

        # Build structured output
        sections = []
        sections.append("SUPPLEMENTARY MATERIALS")
        sections.append("=" * 80)
        sections.append("")

        for title, content in materials.items():
            sections.append(f"## {title}")
            sections.append("-" * 80)
            sections.append(content.strip())
            sections.append("")

        full_text = "\n".join(sections)

        # Apply compression if needed
        if max_chars and len(full_text) > max_chars:
            if compression_strategy == "truncate":
                return self._truncate_materials(full_text, max_chars)
            elif compression_strategy == "summarize":
                return self._summarize_materials(materials, max_chars)
            elif compression_strategy == "smart":
                return self._smart_compress_materials(materials, max_chars)

        return full_text

    def _truncate_materials(self, text: str, max_chars: int) -> str:
        """Simple truncation with ellipsis."""
        if len(text) <= max_chars:
            return text

        truncate_point = max_chars - 100  # Reserve space for message
        return text[:truncate_point] + f"\n\n[... Truncated {len(text) - truncate_point} characters ...]"

    def _summarize_materials(self, materials: Dict[str, str], max_chars: int) -> str:
        """
        Summarize materials by keeping key sections.
        Intelligent extraction of important content.
        """
        # For now, use smart compression
        return self._smart_compress_materials(materials, max_chars)

    def _smart_compress_materials(self, materials: Dict[str, str], max_chars: int) -> str:
        """
        Smart compression that preserves structure and key content.
        Proportionally allocates space to each material.
        """
        sections = []
        sections.append("SUPPLEMENTARY MATERIALS (Compressed)")
        sections.append("=" * 80)
        sections.append("")

        # Calculate overhead
        header_overhead = len("\n".join(sections))
        material_overhead_per_item = 100  # For title, separators, etc.
        total_overhead = header_overhead + (len(materials) * material_overhead_per_item)

        available_chars = max_chars - total_overhead
        if available_chars < 1000:
            # Not enough space, just provide titles
            sections.append("Materials available (content truncated due to space):")
            for title in materials.keys():
                sections.append(f"- {title}")
            return "\n".join(sections)

        # Allocate proportionally based on content size
        total_content_chars = sum(len(content) for content in materials.values())

        for title, content in materials.items():
            # Calculate allocation for this material
            if total_content_chars > 0:
                proportion = len(content) / total_content_chars
                allocated_chars = int(available_chars * proportion)
            else:
                allocated_chars = available_chars // len(materials)

            sections.append(f"## {title}")
            sections.append("-" * 80)

            if len(content) <= allocated_chars:
                sections.append(content.strip())
            else:
                # Keep first part and note truncation
                sections.append(content[:allocated_chars].strip())
                sections.append(f"\n[... {len(content) - allocated_chars} chars truncated ...]")

            sections.append("")

        return "\n".join(sections)

    def get_token_estimate(self, text: str) -> int:
        """
        Estimate token count for text.
        Uses rough heuristic: ~4 chars per token for English text.

        Args:
            text: Text to estimate

        Returns:
            Estimated token count
        """
        return len(text) // 4

    def format_for_model_context(
        self,
        materials: Dict[str, str],
        target_tokens: Optional[int] = None
    ) -> str:
        """
        Format materials for inclusion in model context.

        Args:
            materials: Dictionary of title -> content
            target_tokens: Target token count (None = no limit)

        Returns:
            Formatted string ready for LLM prompt
        """
        if target_tokens:
            max_chars = target_tokens * 4  # Convert tokens to chars
            return self.compact_materials(materials, max_chars, "smart")
        else:
            return self.compact_materials(materials)


def load_imprint_supplementary_materials(
    config: Dict[str, Any],
    use_supplementary: bool = False,
    max_tokens: Optional[int] = None
) -> Optional[str]:
    """
    Load supplementary materials from imprint config.

    Args:
        config: Imprint configuration dictionary
        use_supplementary: Whether to load supplementary materials
        max_tokens: Maximum tokens for supplementary content

    Returns:
        Formatted supplementary materials or None
    """
    if not use_supplementary:
        return None

    supplementary_config = config.get("supplementary_materials")

    if not supplementary_config:
        logger.debug("No supplementary_materials section in config")
        return None

    if not supplementary_config.get("enabled", False):
        logger.debug("Supplementary materials disabled in config")
        return None

    materials_dict = supplementary_config.get("materials", {})

    if not materials_dict:
        logger.debug("No materials defined in supplementary_materials")
        return None

    # Load materials
    loader = SupplementaryMaterialsLoader()
    loaded_materials = loader.load_materials(materials_dict)

    if not loaded_materials:
        logger.warning("No supplementary materials could be loaded")
        return None

    # Format for model context
    formatted = loader.format_for_model_context(loaded_materials, max_tokens)

    logger.info(f"Loaded {len(loaded_materials)} supplementary materials "
                f"({loader.get_token_estimate(formatted)} estimated tokens)")

    return formatted


def get_largest_context_model(model_name: str) -> int:
    """
    Get the largest available context window for a model.

    Args:
        model_name: Model identifier (e.g., "anthropic/claude-sonnet-4-5-20250929")

    Returns:
        Maximum context window in tokens
    """
    # Model context window sizes (in tokens)
    # These are approximate and should be updated as models evolve
    MODEL_CONTEXTS = {
        # Anthropic
        "anthropic/claude-sonnet-4-5-20250929": 200000,
        "anthropic/claude-sonnet-4": 200000,
        "anthropic/claude-opus-4": 200000,
        "anthropic/claude-3.5-sonnet": 200000,

        # OpenAI
        "openai/gpt-4-turbo": 128000,
        "openai/gpt-4": 8192,
        "openai/gpt-3.5-turbo": 16385,

        # Google
        "gemini/gemini-2.5-flash": 1000000,
        "gemini/gemini-2.5-pro": 2000000,
        "gemini/gemini-1.5-flash": 1000000,
        "gemini/gemini-1.5-pro": 2000000,

        # xAI
        "xai/grok-3-latest": 131072,

        # Default
        "default": 8192
    }

    # Normalize model name
    model_key = model_name.lower()

    # Try exact match
    if model_key in MODEL_CONTEXTS:
        return MODEL_CONTEXTS[model_key]

    # Try partial match
    for key, context_size in MODEL_CONTEXTS.items():
        if key in model_key or model_key in key:
            return context_size

    # Default fallback
    logger.warning(f"Unknown model context size for {model_name}, using default")
    return MODEL_CONTEXTS["default"]
