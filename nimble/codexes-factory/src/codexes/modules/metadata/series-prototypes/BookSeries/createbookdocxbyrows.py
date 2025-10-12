'''
input: one book idea with associated characters, locations, events, themes, mcguffins, plot twists, title requirements
output: ordered list of prompts to generate full text of book
'''
import os
import uuid

import pandas as pd
import requests
from docx import Document
from docx.shared import Inches
from openai import Image

from app.utilities.gpt3complete import chatcomplete


def create_book_outline(book_idea, characters, locations, events, themes, mcguffins, plot_twists, title_requirements):
    book_outline = []
    book_outline.append(book_idea)
    book_outline.append(characters)
    book_outline.append(locations)
    book_outline.append(events)
    book_outline.append(themes)
    book_outline.append(mcguffins)
    book_outline.append(plot_twists)
    book_outline.append(title_requirements)
    return book_outline


'''example outline inputs
an intelligent tank receives a message from a distant planet
characters: a tank, a human, a robot
locations: a distant planet, a spaceship
events: a message is received, a message is sent
themes: communication, intelligence
mcguffins: a message
plot twists: the message is from a distant planet
title requirements: genre, wordcount, audience, alignment with series theme
'''


# create text, image, music, video prompts for each pages
# run prompts
# combine pages into book

def create_rows_for_picturebook(text, n=1, size='256x256'):
    rows = []
    text_row, image_row, prompt_row = {}, {}, {}
    thisbookstyle = '''You are MWeltKelly, the staff artist for Nimble Books LLC.
    Your work is to create illustrations for the children's book "I Hate Daddy's Phone."
    Your  work is creative, funny, smart, and engaging.
    Your work is suitable for children ages 3-8.
    Your work is suitable for children who are learning to read.
    Your work should be color cartoon illustrations suitable for children's books.
    Imagine that your style is the visual equivalent of PG Wodehouse.
    Do not use text.  Do not use speech bubbles.  Do not use captions.'''

    for line in text.splitlines():
        text_row = {"string": line, "style": "Body Text"}
        prompt = chatcomplete("ImagePromptGenerator", line, "gpt-3.5-turbo", '37')
        prompt_text = prompt[0]['choices'][0]['message']['content']
        submit_prompt = thisbookstyle + '\n' + prompt_text

        # set a directory to save DALL-E images to
        image_dir_name = "output/picturebooks"
        image_dir = os.path.join(os.curdir, image_dir_name)

        # create the directory if it doesn't yet exist
        if not os.path.isdir(image_dir):
            os.mkdir(image_dir)

        # print the directory to save to
        print(f"{image_dir=}")


        image = Image.create(
            prompt=submit_prompt,
            n=n,
            size=size,
        )
        print(image)

        uuid_string = str(uuid.uuid4())
        generated_image_name = uuid_string + ".png"  # any name you like; the filetype should be .png
        generated_image_filepath = os.path.join(image_dir, generated_image_name)
        image_url = image["data"][0]["url"]  # extract image URL from response
        generated_image = requests.get(image_url).content  # download the image

        with open(generated_image_filepath, "wb") as image_file:
            image_file.write(generated_image)  # write the image to the file

        image_row = {"string": generated_image_filepath, "style": "Image"}
        prompt_row = {"string": prompt_text, "style": "Caption"}
        rows.append(text_row)
        rows.append(image_row)
        rows.append(prompt_row)
        rows.append({"string": "Page Break", "style": "Page Break"})
    return rows

def rows2docx(rows, thisdoc_dir):
    # read list of rows into a dataframe

    df = pd.DataFrame(rows)
    print('df', df)
    # create a list of dicts for each row
    document = Document("resources/docx/FrontMatter2.docx")
    try:
        count = 0
        for index, row in df.iterrows():
            count += 1
            for line in row["string"].splitlines():
                if row["style"] == "Page Break":
                    document.add_page_break()
                    continue
                elif row["style"] == "Image":
                    print('adding picture', line)
                    document.add_picture(line, width=Inches(5.5))
                    continue
                elif row["style"] == "Caption":
                    document.add_paragraph(line, "Caption")
                    continue
                else:
                    document.add_paragraph(line, row["style"])
                document.add_paragraph(line, row["style"])
    except Exception as e:
        errmessage = str(e) + 'error in creating word doc'
        print(errmessage)
    uuid_string = str(uuid.uuid4())
    booktargetpath = thisdoc_dir + uuid_string + '/newbook.docx'
    document.save('output/picturebooks/newbook.docx')

    return df, document

text = '''I have the very bestest daddy.  He's so smart and funny and he plays Planets with me.
He swings me from Earth to Venus to Mercury and back out to Pluto.
But I hate his phone.
'''
text2 = '''
I hate his phone because it's always ringing and he has to answer it.
I hate his phone because it's the newest on the market.
I hate his phone because it's younger than me.
I hate his phone because it's little and dinky and.
I hate his phone because he loves his phone more than he loves me.
I love my Daddy because he's my Daddy.
But I hate his phone.
I am going to destroy his phone.
But first I am going to tell him I love him.
I love you, Daddy.
But if you don't stop playing with your phone and start playing with me, I am going to destroy your phone.
My Daddy hid his phone in his pocket!  Yay!
I love my Daddy.
'''

rows = create_rows_for_picturebook(text=text, n=1, size='256x256')
book = rows2docx(rows, 'output/picturebooks')
